from docx import Document

# Create a new Document
doc = Document()
doc.add_heading('Critical Analysis of Text Classification Algorithms and Techniques', level=1)

# Add the table with headings
table = doc.add_table(rows=1, cols=4)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Technique/Algorithm'
hdr_cells[1].text = 'Application'
hdr_cells[2].text = 'Strengths'
hdr_cells[3].text = 'Limitations'

# Data for each row of the table
rows = [
    ("Bag of Words (BoW)", "Basic text representation", "Simple and easy to implement; effective for small datasets.", "Ignores syntax and word order; poor performance with synonyms and polysemy."),
    ("Thematic Categorizations", "Semantic analysis", "Helps in understanding broader themes; useful for summarization.", "May miss subtle meanings; depends heavily on the quality of thematic definitions."),
    ("Vector Representations (Neural Networks)", "Advanced text representation", "Captures context and semantic meaning; adaptable to different contexts.", "Requires large datasets; computationally intensive."),
    ("Long Short-Term Memory (LSTM)", "Sequence prediction", "Good at capturing long-range dependencies; handles varying-length input sequences well.", "Slower to train; complex model architecture."),
    ("Convolutional Neural Networks (CNNs)", "Feature extraction from text", "Efficient for local pattern recognition (e.g., phrases); parallelizable architecture.", "Less effective for long-range dependencies."),
    ("Recursive Networks", "Hierarchical data processing", "Good for hierarchical problems like parsing; captures nested structure.", "Can be difficult to train; less popular due to difficulty in optimization."),
    ("Support Vector Machine (SVM)", "Binary classification", "Effective in high-dimensional spaces; works well with clear margin of separation.", "Not suitable for larger datasets; requires careful kernel choice."),
    ("Logistic Regression", "Probability estimation", "Simple; outputs probabilities for outcomes; good interpretability.", "Assumes linear relationships; not suitable for complex relationships."),
    ("Softmax Function", "Multiclass classification", "Directly generalizes logistic regression to multiple classes.", "Assumes independence among classes; can be computationally expensive for many classes."),
    ("TF-IDF", "Weighted text representation", "Emphasizes important words; reduces weight of common words across documents.", "Ignores semantics and context; purely statistical approach."),
    ("CNN-LSTM", "Combined feature and sequence learning", "Combines local features extraction with sequence modeling; robust to sequence and spatial inputs.", "Complex to implement and tune; computationally intensive."),
    ("Recurrent Neural Networks (RNNs)", "Sequence prediction", "Can process sequences of data; useful for text where order matters.", "Prone to vanishing and exploding gradient problems; challenging to train."),
    ("FastCNNs", "Efficient feature extraction", "Faster than traditional CNNs; suitable for real-time applications.", "May sacrifice some accuracy for speed; less effective for very complex patterns."),
    ("Bayesian Naïve Bayes Classifiers", "Probabilistic classification", "Incorporates prior knowledge; good for small datasets.", "Assumes independence among features; poor with correlated features."),
    ("Absolute Deviation Factor-based methods (ADF-CTF, ADF-CDF, ADF-CTDF)", "Feature selection", "Reduces dimensionality by selecting features based on their deviation.", "May discard useful information; effectiveness depends on the distribution of data."),
    ("Binary Poor and Rich Optimization Algorithm (HBPRO)", "Feature selection and optimization", "Optimizes feature selection in binary classification tasks; potentially improves accuracy.", "Complex to understand and implement; may be computationally expensive.")
]

# Fill table with data
for tech, app, strengths, limitations in rows:
    row_cells = table.add_row().cells
    row_cells[0].text = tech
    row_cells[1].text = app
    row_cells[2].text = strengths
    row_cells[3].text = limitations

# Save the document
file_path = 'Text_Classification_Algorithms_Analysis.docx'
doc.save(file_path)

file_path
